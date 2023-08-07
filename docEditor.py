from docx import Document
# from docx import Inches
import re
import os
from python_docx_replace import docx_replace
import docxedit
from docx.shared import Pt
import docx2pdf

# class init function will call function to separate strings 
# and run function will call function to replace strings
# and than it will make folder to save new resume.
# Error - while reading the resume file, it is not reading the whole professional experience


class DocEditor:


    def __init__(self, updatedResume):

        resume_path = "resume/ShivamSharmaCV.docx"
        self.document = Document(resume_path)
        self.new_resume = updatedResume
        self.table = self.document.tables[0]
        self.professional_experience = ""
        self.programming_skills = ""
        self.data_frameworks = ""
        self.application_software = ""
        self.tools_and_libraries = ""
        self.main_folder = "curated_resume"

        # checks if folder for curated resume exists, if not then create it.
        self.check_curated_resume_folder(self.main_folder)
        # separate text from the string to be used in the resume.
        try:
            self.separate_text()
        except:
            print("Error in separating text from the string to be used in the resume.")

    def check_curated_resume_folder(self, folder_path):
        # This checks if fodler for curated resume exists or not.
        if os.path.isdir(folder_path):
            pass
        else:
            os.mkdir(folder_path)

    def separate_text(self):
        # This function will separate the text from the string to be used in the resume.
        skills = "skills"
        prog_lang_len = len("Programming Languages") + 2
        data_frameworks_len = len("Data Frameworks") + 2
        application_software_len = len("Application Software") + 2
        tools_and_libraries_len = len("Tools and Libraries") + 2
        
        professional_experience = "professional experience"
        temp_resume = self.new_resume.lower()

        # separate indices
        skills_index = temp_resume.index(skills)
        prof_exp_index = temp_resume.index(professional_experience)
        
        # get professional experience
        self.professional_experience = self.new_resume[prof_exp_index:skills_index].strip().split("\n")[4:]
        self.professional_experience = "\n".join(self.professional_experience)
        # print(self.professional_experience)
        
        # get skills
        all_skills = self.new_resume[skills_index:].strip()
        all_skills_lower = all_skills.lower()

        # get programming skills indices
        programming_skills_index = all_skills_lower.index("programming languages") + prog_lang_len
        data_frameworks_index = all_skills_lower.index("data frameworks") + data_frameworks_len
        application_software_index = all_skills_lower.index("application software") + application_software_len
        tools_and_libraries_index = all_skills_lower.index("tools and libraries") + tools_and_libraries_len
        
        # get programming skills
        self.programming_skills = all_skills[programming_skills_index:(data_frameworks_index-data_frameworks_len)].strip()
        self.data_frameworks = all_skills[data_frameworks_index:(application_software_index-application_software_len)].strip()
        self.application_software = all_skills[application_software_index:(tools_and_libraries_index-tools_and_libraries_len)].strip()
        self.tools_and_libraries = all_skills[tools_and_libraries_index:].strip().split("\n")[0]

        # print(self.tools_and_libraries)
        
    def edit_work_experience(self, data):
        # This function edits the work experience section of the resume.
        experience_cell =  self.table.cell(2, 0)
        experience_cell.text = data
        experience_cell.paragraphs[0].runs[0].font.size = Pt(9)
        experience_cell.paragraphs[0].runs[0].font.name = "Calibri Light" 
    
    def edit_skills_section(self):
        # This function edits the skills section of the resume.
        programming_languages_cell = self.table.rows[4].cells[1]
        data_frameworks_cell = self.table.rows[5].cells[1]
        application_software_cell = self.table.rows[6].cells[1]
        tools_and_libraries_cell = self.table.rows[7].cells[1]
        # Changing programming language
        programming_languages_cell.text = self.programming_skills
        programming_languages_cell.paragraphs[0].runs[0].font.size = Pt(9)
        programming_languages_cell.paragraphs[0].runs[0].font.name = "Calibri Light"
        # Changing data frameworks
        data_frameworks_cell.text = self.data_frameworks
        data_frameworks_cell.paragraphs[0].runs[0].font.size = Pt(9)
        data_frameworks_cell.paragraphs[0].runs[0].font.name = "Calibri Light"
        # # Changing application software
        application_software_cell.text = self.application_software
        application_software_cell.paragraphs[0].runs[0].font.size = Pt(9)
        application_software_cell.paragraphs[0].runs[0].font.name = "Calibri Light"
        # # Changing tools and libraries
        tools_and_libraries_cell.text = self.tools_and_libraries
        tools_and_libraries_cell.paragraphs[0].runs[0].font.size = Pt(9)
        tools_and_libraries_cell.paragraphs[0].runs[0].font.name = "Calibri Light"
        
    def run(self, company_name="comapny_name"):
        # This function will run the whole process of editing the resume.
        # edit work experience section of the resume.
        try:
            self.edit_work_experience(self.professional_experience)
        except:
            print("Error in editing work experience section of the resume.")
        # pass
        # # edit skills section of the resume.
        try:
            self.edit_skills_section()
        except:
            print("Error in editing skills section of the resume.")
        
        self.check_curated_resume_folder(os.path.join(self.main_folder, company_name))
        self.document.save(os.path.join(self.main_folder, company_name, "ShivamSharmaCV.docx"))
        docx2pdf.convert(os.path.join(self.main_folder, company_name, "ShivamSharmaCV.docx"), os.path.join(self.main_folder, company_name, "ShivamSharmaCV.pdf"))


if __name__ == '__main__':
    resume = """Shivam Sharma"""
    doc = DocEditor(resume)
    doc.run("Acti")
